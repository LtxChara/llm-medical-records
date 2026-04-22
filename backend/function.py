import asyncio
import logging
import json
from typing import Dict, Any, Tuple, Optional

from openai import OpenAI, RateLimitError, APIError, APIConnectionError
from med_prompt import system_prompt_templates, user_prompt_templates
from chat_memory import MedicalHistoryManager
from schemas import FIELD_SCHEMA_MAP
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO

from config import settings

logger = logging.getLogger(__name__)

# Order of fields to process
FIELDS_ORDER = ["主诉", "现病史", "既往史", "过敏史", "诊断"]
# Maximum number of supplement rounds per field before forcing continuation
MAX_SUPPLEMENT_PER_FIELD = 3


class SessionStore:
    """Async-safe session store with per-session locks, TTL eviction, and background cleanup."""

    def __init__(self, ttl_seconds: int, max_sessions: int = 1000):
        self._data: Dict[str, MedicalHistoryManager] = {}
        self._locks: Dict[str, asyncio.Lock] = {}
        self._last_access: Dict[str, float] = {}
        self.ttl_seconds = ttl_seconds
        self._max_sessions = max_sessions
        self._cleanup_task: Optional[asyncio.Task] = None
        self._cleanup_interval = 600  # 每10分钟执行一次后台清理

    def _now(self) -> float:
        return asyncio.get_event_loop().time()

    async def get_or_create(self, session_id: str) -> Tuple[MedicalHistoryManager, asyncio.Lock]:
        now = self._now()
        expired = [
            sid for sid, t in self._last_access.items()
            if now - t > self.ttl_seconds
        ]
        for sid in expired:
            self._data.pop(sid, None)
            self._locks.pop(sid, None)
            self._last_access.pop(sid, None)
            logger.info(f"Expired session removed: {sid}")

        if session_id not in self._data:
            self._data[session_id] = MedicalHistoryManager()
            self._locks[session_id] = asyncio.Lock()
        self._last_access[session_id] = now
        return self._data[session_id], self._locks[session_id]

    def exists(self, session_id: str) -> bool:
        return session_id in self._data

    async def _cleanup_loop(self):
        """后台定时清理循环，防止废弃会话永久驻留内存。"""
        while True:
            await asyncio.sleep(self._cleanup_interval)
            try:
                await self._do_cleanup()
            except Exception as e:
                logger.exception(f"Background session cleanup failed: {e}")

    async def _do_cleanup(self):
        """执行TTL过期清理与LRU上限淘汰。"""
        now = self._now()
        expired = [
            sid for sid, t in self._last_access.items()
            if now - t > self.ttl_seconds
        ]
        for sid in expired:
            self._data.pop(sid, None)
            self._locks.pop(sid, None)
            self._last_access.pop(sid, None)
        if expired:
            logger.info(f"Background cleanup removed {len(expired)} expired sessions.")

        # LRU淘汰：如果会话数超过上限，淘汰最久未访问的
        if len(self._data) > self._max_sessions:
            sorted_sessions = sorted(self._last_access.items(), key=lambda x: x[1])
            to_evict = sorted_sessions[:len(self._data) - self._max_sessions]
            for sid, _ in to_evict:
                self._data.pop(sid, None)
                self._locks.pop(sid, None)
                self._last_access.pop(sid, None)
            logger.warning(f"LRU eviction removed {len(to_evict)} sessions (limit: {self._max_sessions}).")

    def start_cleanup(self):
        """启动后台清理任务。应在事件循环启动后调用（如FastAPI startup事件）。"""
        if self._cleanup_task is None or self._cleanup_task.done():
            self._cleanup_task = asyncio.create_task(self._cleanup_loop())
            logger.info("SessionStore background cleanup task started.")

    def stop_cleanup(self):
        """停止后台清理任务。应在应用关闭时调用（如FastAPI shutdown事件）。"""
        if self._cleanup_task and not self._cleanup_task.done():
            self._cleanup_task.cancel()
            logger.info("SessionStore background cleanup task stopped.")


store = SessionStore(ttl_seconds=settings.session_ttl_seconds, max_sessions=1000)


def _sanitize_user_input(content: str) -> str:
    """
    对用户输入进行基础清理，防御Prompt注入与异常字符。
    - 去除不可打印的控制字符（保留换行、制表符）。
    - 限制单条输入最大长度，防止超大输入导致Token爆炸或内存占用。
    """
    if not content:
        return ""
    content = ''.join(ch for ch in content if ch.isprintable() or ch in '\n\r\t')
    max_len = 8000
    if len(content) > max_len:
        logger.warning(f"User input truncated from {len(content)} to {max_len} characters.")
        content = content[:max_len]
    return content


try:
    client = OpenAI(
        api_key=settings.api_key,
        base_url=settings.base_url,
        timeout=settings.llm_timeout,
    )
    logger.info(
        f"OpenAI client initialized for model: {settings.llm_model_name} "
        f"at {settings.base_url}"
    )
except Exception as e:
    logger.critical(f"Failed to initialize OpenAI client: {e}")
    client = None

# --- Core LLM Interaction Logic ---

async def call_llm(
    field: str,
    history: MedicalHistoryManager
) -> Tuple[Optional[Dict[str, Any]], Optional[str]]:
    """
    Calls the LLM for a specific field using the complete conversation history.

    关键改进：
    - 使用 history.to_openai_format() 获取完整对话历史（System + Human + AI），
      而非仅拼接用户输入，从而保留多轮补充的上下文。
    - 仅将第一条用户消息包装在 user_prompt_template 中，后续补充消息保持原样。
    - 系统提示中注入 JSON Schema，提升结构化输出质量。

    Args:
        field: The medical field to query (e.g., "主诉").
        history: The MedicalHistoryManager instance for the current session.

    Returns:
        A tuple: (parsed_json_response, error_message).
        If successful, parsed_json_response is the dict, error_message is None.
        If failed, parsed_json_response is None, error_message contains the error detail.
    """
    if not client:
        return None, "LLM client not initialized."
    if field not in history.histories:
        return None, f"History type '{field}' not found in history manager."

    schema_cls = FIELD_SCHEMA_MAP.get(field)
    user_prompt_template = user_prompt_templates.get(field)
    system_prompt = system_prompt_templates.get(field)

    if not schema_cls or not user_prompt_template or not system_prompt:
        return None, f"Prompt template or schema missing for field: {field}"

    # 构建增强版系统提示，注入JSON Schema以提升结构化输出质量
    schema_json = schema_cls.model_json_schema()
    enhanced_system_prompt = (
        f"{system_prompt}\n\n"
        f"你必须严格按照以下 JSON Schema 输出，不要添加任何其他说明文字：\n"
        f"{json.dumps(schema_json, ensure_ascii=False, indent=2)}"
    )

    # 获取该字段的完整对话历史（OpenAI格式）
    try:
        historical_messages = history.to_openai_format(field)
    except ValueError as e:
        return None, str(e)

    messages = []

    # 处理系统提示：始终将增强版系统提示置于首位
    if historical_messages and historical_messages[0]["role"] == "system":
        messages.append({"role": "system", "content": enhanced_system_prompt})
        # 移除原始系统提示，避免重复
        historical_messages = historical_messages[1:]
    else:
        messages.append({"role": "system", "content": enhanced_system_prompt})

    # 处理历史消息：仅将第一条用户消息包装在user_prompt_template中，
    # 后续补充消息保持原样，以保留多轮对话上下文
    first_user_processed = False
    for msg in historical_messages:
        if msg["role"] == "user" and not first_user_processed:
            # 包装第一条用户消息（通常是初始患者描述）
            wrapped_content = user_prompt_template + msg["content"] + "'''"
            messages.append({"role": "user", "content": wrapped_content})
            first_user_processed = True
        else:
            messages.append(msg)

    # 如果历史中没有用户消息，尝试从所有用户输入构建一条
    if not first_user_processed:
        user_input_history = history.get_user_all_input(field)
        if not user_input_history.strip():
            logger.warning(f"No user input history found for field '{field}'. Skipping LLM call.")
            return {field: ""}, None
        wrapped_content = user_prompt_template + user_input_history + "'''"
        messages.append({"role": "user", "content": wrapped_content})

    logger.debug(f"Calling LLM for field '{field}' with messages: {messages}")

    max_retries = settings.llm_max_retries
    base_delay = 1  # Base delay in seconds for retries

    for attempt in range(max_retries + 1):
        try:
            response = await asyncio.to_thread(
                client.chat.completions.create,
                model=settings.llm_model_name,
                messages=messages,
                temperature=0.5,
                top_p=0.9,
                response_format={"type": "json_object"},
            )
            content = response.choices[0].message.content
            logger.debug(f"LLM raw response for '{field}': {content}")

            if not content:
                logger.warning(f"LLM returned empty content for field '{field}'.")
                return {field: ""}, None

            try:
                validated = schema_cls.model_validate_json(content)
                return validated.model_dump(), None
            except Exception as parse_e:
                logger.error(
                    f"Failed to parse/repair LLM response JSON for '{field}'. "
                    f"Error: {parse_e}. Raw content: '{content}'"
                )
                return None, f"Failed to parse LLM response: {parse_e}"

        except RateLimitError as rle:
            logger.warning(f"Rate limit hit calling LLM for '{field}'. Attempt {attempt + 1}/{max_retries + 1}. Error: {rle}")
            if attempt >= max_retries:
                return None, f"Rate limit exceeded after {max_retries + 1} attempts."
            await asyncio.sleep(base_delay * (2 ** attempt))
        except (APIError, APIConnectionError) as apie:
            logger.error(f"API error calling LLM for '{field}'. Attempt {attempt + 1}/{max_retries + 1}. Error: {apie}")
            if attempt >= max_retries:
                return None, f"API error after {max_retries + 1} attempts: {apie}"
            await asyncio.sleep(base_delay * (2 ** attempt))
        except Exception as e:
            logger.exception(f"Unexpected error calling LLM for '{field}'. Attempt {attempt + 1}/{max_retries + 1}. Error: {e}")
            if attempt >= max_retries:
                return None, f"Unexpected error after {max_retries + 1} attempts: {e}"
            await asyncio.sleep(base_delay)

    return None, f"LLM call failed for field '{field}' after {max_retries + 1} attempts."


async def getLLMReturn(
    session_id: str,
    initial_text: Optional[str] = None,
    history_manager: Optional[MedicalHistoryManager] = None,
    target_field: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Orchestrates the process of calling the LLM for required fields,
    handles the user supplementation workflow.

    关键改进：
    1. 初始文本同时注入所有字段的历史记录，确保每个字段都有原始患者描述作为提取语料。
    2. 使用 history_manager 的 completed_fields 跟踪字段完成状态，
       避免已完成的字段被重复调用，提升效率与稳定性。
    3. 补充场景（target_field）只重新处理目标字段及受影响的后续字段（诊断），
       而非重跑全部字段。
    4. 诊断字段独立注入已提取字段的聚合上下文，消除与原始Prompt的语义冲突。

    Args:
        session_id: The unique identifier for the user session.
        initial_text: The initial user input (only provided on the first call via /process).
        history_manager: The MedicalHistoryManager for this session.
        target_field: 指定需要重新处理的字段（用于 /supplement 场景）。

    Returns:
        A dictionary representing the API response:
        - {"status": "success", "result": {...}} if all fields complete.
        - {"status": "incomplete", "field": "...", "message": "..."} if supplementation needed.
        - {"status": "error", "message": "..."} if a critical error occurs.
    """
    if not isinstance(session_id, str) or not session_id.strip():
        return {"status": "error", "message": "Session ID cannot be empty."}

    # Get or create history manager for the session
    if history_manager is None:
        history_manager, _ = await store.get_or_create(session_id)
        logger.info(f"Creating new MedicalHistoryManager for session_id: {session_id}")

    # 如果提供了初始文本，将其注入所有字段的历史记录中
    # 确保每个字段都有原始患者描述作为提取语料
    if initial_text and isinstance(initial_text, str) and initial_text.strip():
        for field in FIELDS_ORDER:
            try:
                history_manager.add_message(field, "user", initial_text.strip())
            except (ValueError, IndexError) as e:
                logger.error(f"Failed to add initial text to '{field}' for session {session_id}: {e}")
        logger.info(f"Added initial text to all fields for session {session_id}.")

    # 确定本次需要处理的字段列表
    if target_field:
        if target_field not in FIELDS_ORDER:
            return {"status": "error", "message": f"Invalid target field: '{target_field}'."}
        target_idx = FIELDS_ORDER.index(target_field)
        fields_to_process = [target_field]
        # 补充场景：处理目标字段及其之后所有尚未完成的字段
        # 确保流程能从前一个断点继续推进到诊断
        for f in FIELDS_ORDER[target_idx + 1:]:
            if not history_manager.is_field_completed(f):
                fields_to_process.append(f)
        # 如果诊断之前已完成，且目标字段在诊断之前，也需要重新评估诊断
        if target_field != "诊断" and history_manager.is_field_completed("诊断"):
            history_manager.reset_field_status("诊断")
            if "诊断" not in fields_to_process:
                fields_to_process.append("诊断")
            logger.info(f"Diagnosis will be re-evaluated due to supplement in '{target_field}'.")
    else:
        # 初始场景：处理所有尚未完成的字段
        fields_to_process = [f for f in FIELDS_ORDER if not history_manager.is_field_completed(f)]

    # 如果所有字段都已完成，直接返回缓存结果
    if not fields_to_process:
        logger.info(f"All fields already completed for session {session_id}.")
        return {
            "status": "success",
            "result": {f: history_manager.get_field_result(f) for f in FIELDS_ORDER},
            "message": "处理成功",
        }

    # 定义内部辅助函数，处理单个字段的结果并更新历史与状态
    def _handle_field_result(field: str, llm_result: Optional[Dict[str, Any]], error_msg: Optional[str]):
        if error_msg:
            logger.error(f"Error processing field '{field}' for session {session_id}: {error_msg}")
            return "error", {"status": "error", "message": f"Failed to process field '{field}': {error_msg}"}

        if llm_result is None:
            logger.error(f"LLM result was None but no error message for field '{field}' (session {session_id}).")
            return "error", {"status": "error", "message": f"Internal error processing field '{field}'."}

        field_value = llm_result.get(field, "")
        needs_supplement = bool(llm_result.get("needs_supplement", False))
        supplement_question = llm_result.get("supplement_question") or field_value

        if needs_supplement:
            supplement_count = history_manager.get_supplement_count(field)
            if supplement_count >= MAX_SUPPLEMENT_PER_FIELD:
                logger.warning(
                    f"Max supplement reached for field '{field}' in session {session_id} "
                    f"(count={supplement_count}). Forcing continuation."
                )
                # 强制使用该字段结果，标记为完成（即使内容可能不够理想）
                forced_value = field_value if field_value and field_value != "信息不足" else ""
                history_manager.set_field_result(field, forced_value)
                history_manager.mark_field_completed(field)
                try:
                    history_manager.add_message(field, "ai", f"[已达最大补充次数，强制使用当前结果] {supplement_question}")
                except ValueError as e:
                    logger.error(f"Failed to add AI message to history for field '{field}' (session {session_id}): {e}")
                return "success", None

            logger.info(f"Supplementation needed for field '{field}' in session {session_id}. Message: {supplement_question}")
            try:
                history_manager.add_message(field, "ai", supplement_question)
            except ValueError as e:
                logger.error(f"Failed to add AI supplement message to history for field '{field}' (session {session_id}): {e}")
            return "incomplete", {
                "status": "incomplete",
                "field": field,
                "message": supplement_question,
            }
        else:
            # 字段提取成功，持久化结果并标记完成
            history_manager.set_field_result(field, field_value)
            history_manager.mark_field_completed(field)
            logger.debug(f"Field '{field}' processed successfully for session {session_id}. Value: {field_value}")
            try:
                ai_response_content = json.dumps(field_value, ensure_ascii=False) if not isinstance(field_value, str) else field_value
                history_manager.add_message(field, "ai", ai_response_content)
            except ValueError as e:
                logger.error(f"Failed to add AI message to history for field '{field}' (session {session_id}): {e}")
            return "success", None

    # --- Phase 1: Concurrent extraction of independent fields ---
    independent_fields = [f for f in fields_to_process if f != "诊断"]
    if independent_fields:
        logger.info(f"Phase 1: Concurrently processing fields {independent_fields} for session {session_id}...")

        tasks = [call_llm(field, history_manager) for field in independent_fields]
        independent_results = await asyncio.gather(*tasks, return_exceptions=True)

        for field, result in zip(independent_fields, independent_results):
            if isinstance(result, Exception):
                logger.exception(f"Unexpected exception during concurrent processing of field '{field}' for session {session_id}: {result}")
                return {"status": "error", "message": f"Unexpected error processing field '{field}': {result}"}

            llm_result, error_msg = result
            status, payload = _handle_field_result(field, llm_result, error_msg)
            if status == "error":
                return payload
            elif status == "incomplete":
                return payload

    # --- Phase 2: Sequential diagnosis (depends on previous fields) ---
    if "诊断" in fields_to_process:
        diagnosis_field = "诊断"
        logger.info(f"Phase 2: Processing field '{diagnosis_field}' for session {session_id}...")

        # 注入已提取字段的聚合上下文到诊断历史
        # 仅当诊断历史中没有该上下文消息时才添加，避免重复注入
        context_lines = []
        for f in FIELDS_ORDER[:-1]:  # 除诊断外的字段
            if history_manager.is_field_completed(f):
                val = history_manager.get_field_result(f)
                context_lines.append(f"【{f}】{val}")

        if context_lines:
            aggregated_context = "\n".join(context_lines)
            context_message = (
                f"以下为患者已提取的医疗记录信息，请据此进行诊断：\n{aggregated_context}"
            )
            # 检查是否已存在相同上下文消息，避免重复注入
            existing_input = history_manager.get_user_all_input(diagnosis_field)
            if context_message not in existing_input:
                try:
                    history_manager.add_message(diagnosis_field, "user", context_message)
                    logger.info(f"Injected aggregated context into '{diagnosis_field}' for session {session_id}.")
                except ValueError as e:
                    logger.error(f"Failed to inject diagnosis context for session {session_id}: {e}")

        llm_result, error_msg = await call_llm(diagnosis_field, history_manager)
        status, payload = _handle_field_result(diagnosis_field, llm_result, error_msg)
        if status == "error":
            return payload
        elif status == "incomplete":
            return payload

    # --- Final Response ---
    logger.info(f"All fields processed successfully for session {session_id}.")
    final_result = {f: history_manager.get_field_result(f) for f in FIELDS_ORDER}
    return {
        "status": "success",
        "result": final_result,
        "message": "处理成功",
    }


# --- Document Generation ---

def generate_word_document(params: Dict[str, Any]) -> BytesIO:
    """Generates a formatted Word document from structured patient data."""
    logger.info("Starting Word document generation.")
    try:
        # Validate input structure
        patient_info = params.get('patient_info', {})
        visit_info = params.get('visit_info', {})
        medical_content = params.get('medical_content', {})

        if not patient_info or not visit_info or not medical_content:
            logger.warning("Missing required data sections (patient_info, visit_info, medical_content) for document generation.")
            raise ValueError("Missing required data for document generation.")

        doc = Document()
        # Set default font for Chinese characters
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        # Title
        title = doc.add_heading('中山三院康复科医疗报告', level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Spacing after title
        doc.add_paragraph()

        # Patient Info Section
        doc.add_heading('患者基本信息', level=1).style.font.size = Pt(14)
        patient_mapping = {
            'name': '姓名', 'gender': '性别', 'age': '年龄',
            'id_card': '身份证号', 'phone': '电话号码', 'address': '家庭住址'
        }
        patient_data = []
        for key, label in patient_mapping.items():
            value = patient_info.get(key, '未提供')
            patient_data.append(f"{label}: {value}")
        # Format patient info (e.g., two items per line)
        for i in range(0, len(patient_data), 2):
            line = '    '.join(patient_data[i:i+2])
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)

        # Visit Info Section
        doc.add_heading('就诊信息', level=1).style.font.size = Pt(14)
        visit_mapping = {
            'visit_date': '就诊日期', 'department': '科室', 'doctor': '医生姓名'
        }
        visit_data = []
        for key, label in visit_mapping.items():
            value = visit_info.get(key, '未提供')
            visit_data.append(f"{label}: {value}")
        for i in range(0, len(visit_data), 2):
            line = '    '.join(visit_data[i:i+2])
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)

        # Medical Content Section (Using FIELDS_ORDER for consistency)
        doc.add_heading('医疗记录', level=1).style.font.size = Pt(14)
        medical_display_mapping = {
            "主诉": "主诉", "现病史": "现病史", "既往史": "既往史",
            "过敏史": "过敏史", "诊断": "初步诊断及建议"
        }

        for field in FIELDS_ORDER:
            label = medical_display_mapping.get(field, field)
            value = medical_content.get(field, '未记录')
            # Ensure value is a string for display
            if not isinstance(value, str):
                try:
                    value_str = json.dumps(value, ensure_ascii=False)
                except Exception:
                    value_str = str(value)
            else:
                value_str = value

            doc.add_paragraph(f"{label}:", style='Intense Quote')
            doc.add_paragraph(value_str if value_str else "无")
            doc.add_paragraph()

        # Save to buffer
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        logger.info("Word document generated successfully.")
        return docx_buffer

    except ValueError as ve:
        logger.error(f"Value error during document generation: {ve}")
        raise
    except Exception as e:
        logger.exception("Unexpected error during Word document generation.")
        raise
